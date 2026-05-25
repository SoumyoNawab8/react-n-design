from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys

doc = Document('/Users/macworld/Downloads/Copy of Soumyajyoti Majumdar.docx')

# Helper to find paragraph index containing text
def find_para_index(text):
    for i, p in enumerate(doc.paragraphs):
        if text.lower() in p.text.lower():
            return i
    return -1

# 1. Update Primary Skills to be React-heavy
idx = find_para_index("Primary Skills")
if idx >= 0:
    para = doc.paragraphs[idx]
    para.clear()
    run = para.add_run("Primary Skills: React.js (Expert), Next.js, TypeScript, Node.js, Nuxt.js, AWS, MySQL, Pinia, FastAPI, Redux, Claude Code, GenAI Integration, Component Library Architecture")
    run.bold = True
    run.font.size = Pt(11)
    print("Updated Primary Skills")

# 2. Update Summary to mention react-n-design
idx = find_para_index("Summary")
if idx >= 0:
    summary_idx = idx + 1
    if summary_idx < len(doc.paragraphs):
        summary_para = doc.paragraphs[summary_idx]
        new_text = summary_para.text.strip()
        if new_text and not new_text.endswith('.'):
            new_text += '.'
        addition = " Passionate about open-source development and creator of react-n-design (npm package) — an AI-native, Neomorphic React component library with 40+ accessible components, RSC support, and zero-config installation."
        summary_para.clear()
        run = summary_para.add_run(new_text + addition)
        run.font.size = Pt(10)
        print("Updated Summary")

# 3. Update Key Skills - add React ecosystem and package dev
idx = find_para_index("Key Skills")
if idx >= 0:
    skills_idx = idx + 1
    if skills_idx < len(doc.paragraphs):
        skills_para = doc.paragraphs[skills_idx]
        original = skills_para.text
        new_skills = original + ", React Component Library Development, npm Package Publishing & Distribution, Styled-Components, Framer Motion, Storybook Documentation, Vitest/Jest Testing, Axe-core Accessibility, RSC (React Server Components), Semantic Versioning, Rollup Bundling"
        skills_para.clear()
        run = skills_para.add_run(new_skills)
        run.font.size = Pt(10)
        print("Updated Key Skills")

# 4. Add Open Source section before Experience
idx = find_para_index("Experience")
if idx >= 0:
    exp_para = doc.paragraphs[idx]
    parent = exp_para._element.getparent()
    exp_index = parent.index(exp_para._element)

    # Insert Open Source heading
    os_heading = doc.add_paragraph()
    os_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = os_heading.add_run("Open Source Projects")
    run.bold = True
    run.font.size = Pt(14)
    os_heading._element.getparent().remove(os_heading._element)
    parent.insert(exp_index, os_heading._element)

    # Project title
    proj_title = doc.add_paragraph()
    run = proj_title.add_run("react-n-design (npm package) — AI-native, Neomorphic React Component Library")
    run.bold = True
    run.font.size = Pt(11)
    proj_title._element.getparent().remove(proj_title._element)
    parent.insert(exp_index + 1, proj_title._element)

    # Bullet points
    bullets = [
        "Built and published a production-grade React component library (npm: react-n-design) with 40+ accessible, themeable components",
        "AI-native components: AIChat (streaming chat UI with markdown/typing indicators), CommandPalette (Cmd+K fuzzy search), Calendar, CodeBlock, Markdown renderer",
        "Zero-config installation: npm install react-n-design — no peer dependency friction, styled-components bundled as regular dependency",
        "Full RSC (React Server Components) support with dual entry points (react-n-design & react-n-design/rsc) for Next.js App Router",
        "Accessibility-first: axe-core validated, WCAG 2.2 AA compliance, keyboard navigation, ARIA live regions, prefers-reduced-motion support",
        "Developer experience: 41+ Vitest tests across 10 components, Storybook documentation, TypeScript strict mode, Biome linting, GitHub Actions CI",
        "Advanced features: RTL (right-to-left) support, print styles, touch-optimized CSS, dark/light/system theme modes",
        "Tech stack: React 18+, TypeScript, styled-components v6, framer-motion, rollup (CJS/ESM dual output), vitest + @testing-library/react"
    ]

    for i, bullet in enumerate(bullets):
        bp = doc.add_paragraph("• " + bullet)
        bp.runs[0].font.size = Pt(10)
        bp._element.getparent().remove(bp._element)
        parent.insert(exp_index + 2 + i, bp._element)

    # Spacing
    spacer = doc.add_paragraph()
    spacer._element.getparent().remove(spacer._element)
    parent.insert(exp_index + 2 + len(bullets), spacer._element)

    print("Added Open Source Projects section")

# Save output
output_path = '/Users/macworld/Downloads/Soumyajyoti Majumdar - Updated.docx'
doc.save(output_path)
print(f"\nCV updated and saved to: {output_path}")
